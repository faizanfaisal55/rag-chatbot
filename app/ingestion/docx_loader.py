from docx import Document


def load_docx_file(file_path: str):
    """
    Load a Word (.docx) document.

    Returns a list of "pages" (same shape as PDF loader)
    so it plugs directly into the existing chunk_pages()
    pipeline.

    Word documents don't have a real page concept in the
    file format itself (page breaks are rendering-dependent,
    not stored as fixed boundaries), so the whole document
    is treated as a single page — consistent with how
    text_loader.py handles .txt/.md files.
    """

    document = Document(file_path)

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    # Also extract text from tables, since resumes/reports
    # often use tables for structured info (skills, dates).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    text = "\n".join(paragraphs)

    if not text.strip():
        return []

    return [{
        "page_number": 1,
        "text": text
    }]